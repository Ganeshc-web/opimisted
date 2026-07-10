"""Generate templates/report_template.docx with labelled placeholder variables."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# (label, placeholder) — placeholder must match report_service.py replacements
SECTIONS = {
    "Client Information": [
        ("Client Name", "{{client_name}}"),
        ("Client Age", "{{client_age}}"),
        ("Date of Birth", "{{client_dob}}"),
        ("Occupation", "{{client_occupation}}"),
        ("Company", "{{client_company}}"),
        ("Mobile", "{{mobile}}"),
        ("Email", "{{email}}"),
        ("Address", "{{residential_address}}"),
        ("Spouse Name", "{{spouse_name}}"),
        ("Spouse Age", "{{spouse_age}}"),
        ("Spouse DOB", "{{spouse_dob}}"),
        ("Spouse Occupation", "{{spouse_occupation}}"),
        ("Spouse Company", "{{spouse_company}}"),
        ("Report Date", "{{report_date}}"),
        ("Assessment ID", "{{assessment_id}}"),
    ],
    "Client Retirement Plan": [
        ("Retirement Age", "{{client_retirement_age}}"),
        ("Years to Retirement", "{{client_years_to_retirement}}"),
        ("Monthly Expense Today", "{{client_expense_today}}"),
        ("Monthly Expense at Retirement", "{{client_expense_at_retirement}}"),
        ("Required Corpus", "{{client_corpus}}"),
        ("EPF Corpus at Retirement", "{{client_pf_corpus}}"),
        ("Net Corpus Gap", "{{client_net_corpus}}"),
        ("Monthly SIP Required", "{{client_monthly_sip}}"),
        ("Lump Sum Required Today", "{{client_lump_sum}}"),
    ],
    "Spouse Retirement Plan": [
        ("Retirement Age", "{{spouse_retirement_age}}"),
        ("Years to Retirement", "{{spouse_years_to_retirement}}"),
        ("Monthly Expense Today", "{{spouse_expense_today}}"),
        ("Monthly Expense at Retirement", "{{spouse_expense_at_retirement}}"),
        ("Required Corpus", "{{spouse_corpus}}"),
        ("EPF Corpus at Retirement", "{{spouse_pf_corpus}}"),
        ("Net Corpus Gap", "{{spouse_net_corpus}}"),
        ("Monthly SIP Required", "{{spouse_monthly_sip}}"),
        ("Lump Sum Required Today", "{{spouse_lump_sum}}"),
    ],
    "Financial Goals": [
        ("Goals", "{{goals_table}}"),
        ("Total Monthly SIP (All Goals)", "{{total_goals_monthly_sip}}"),
    ],
    "Insurance": [
        ("Total Insurance Required", "{{total_insurance_required}}"),
    ],
    "Assumptions & Rates": [
        ("Pre-Retirement Inflation", "{{inflation_pre}}"),
        ("Pre-Retirement ROI", "{{roi_pre}}"),
        ("Post-Retirement Inflation", "{{inflation_post}}"),
        ("Post-Retirement ROI", "{{roi_post}}"),
        ("Calculated At", "{{calculated_at}}"),
    ],
}


def add_label_value_row(document: Document, label: str, placeholder: str) -> None:
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(11)
    value_run = paragraph.add_run(placeholder)
    value_run.font.size = Pt(11)
    value_run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def add_section_heading(document: Document, text: str) -> None:
    heading = document.add_heading(text, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)


def build_template() -> Document:
    document = Document()

    title = document.add_heading("Financial Planning Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph("Wealth Wisdom Platform")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle.runs:
        subtitle.runs[0].font.size = Pt(12)
        subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    document.add_paragraph()

    for section_name, rows in SECTIONS.items():
        add_section_heading(document, section_name)
        for label, placeholder in rows:
            if placeholder == "{{goals_table}}":
                document.add_paragraph()
                goals_para = document.add_paragraph()
                goals_run = goals_para.add_run("{{goals_table}}")
                goals_run.font.size = Pt(11)
                goals_run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
            else:
                add_label_value_row(document, label, placeholder)
        document.add_paragraph()

    footer = document.add_paragraph(
        "This report is for informational purposes only and does not constitute "
        "financial advice. © Wealth Wisdom Platform."
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer.runs:
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return document


def main() -> None:
    output_path = Path(__file__).resolve().parent / "templates" / "report_template.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = build_template()
    document.save(output_path)
    print(f"Template created: {output_path}")
    print("Open in Word, add logo/branding, keep {{placeholders}} unchanged.")


if __name__ == "__main__":
    main()
