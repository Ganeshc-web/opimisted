"""Build the dynamic insurance table inside insurance_cal.docx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.core.formatters import fmt_inr_report
from app.core.formulas import (
    build_insurance_rows,
    compute_all_goals,
    monthly_effective_rate,
    real_rate,
)
from app.models.rate_config import RateConfig

INSURANCE_TABLE_MARKER = "{{insurance_table}}"
INSURANCE_HEADER_FILL = "F9A867"
INSURANCE_HEADERS = [
    "Goals to be protected",
    "For/ After years",
    "Amount",
    "Today's/ Future cost",
    "INSURANCE REQUIRED",
]
INSURANCE_COLUMN_RATIOS = (5.8, 2.0, 2.8, 3.0, 3.4)
EMU_PER_INCH = 914400
TWIPS_PER_CM = 567
DEFAULT_TABLE_CLEARANCE_CM = 7.2
EXTRA_TABLE_GAP_CM = 0.25
INSURANCE_HEADER_FONT_PT = 12
INSURANCE_BODY_FONT_PT = 11
INSURANCE_HEADER_ROW_HEIGHT_TWIPS = 850
INSURANCE_DATA_ROW_HEIGHT_TWIPS = 750
INSURANCE_FOOTER_ROW_HEIGHT_TWIPS = 820
INSURANCE_CELL_MARGIN_V_TWIPS = 160
INSURANCE_CELL_MARGIN_H_TWIPS = 120
DISPLAY_NEED_OVERRIDES = {
    "Household Expenses": "Household Expense",
    "Retirement Income (50%)": "Retirement Income(50 %)",
}


def fmt_inr_table(value: float) -> str:
    """Indian currency with a non-breaking space after ₹ (avoids symbol/value wrap)."""
    try:
        amount = float(value or 0)
    except (TypeError, ValueError):
        return ""
    if amount == 0:
        return ""
    formatted = fmt_inr_report(amount)
    if formatted.startswith("₹"):
        # NBSP keeps "₹" and digits on one line in narrow table cells.
        return f"₹\u00a0{formatted[1:].strip()}"
    return formatted


_CHILD_NEED_RE = re.compile(r"^Child\s+(\d+)\s+(.+)$", re.IGNORECASE)


def _first_name(full_name: str | None) -> str:
    parts = (full_name or "").strip().split()
    return parts[0] if parts else ""


def display_need_label(
    need: str,
    child_names_by_number: dict[int, str] | None = None,
) -> str:
    """Human-readable insurance row label; prefers real child names over Child N."""
    if need in DISPLAY_NEED_OVERRIDES:
        return DISPLAY_NEED_OVERRIDES[need]

    match = _CHILD_NEED_RE.match((need or "").strip())
    if not match:
        return need

    child_number = int(match.group(1))
    goal_suffix = match.group(2).strip()
    names = child_names_by_number or {}
    child_name = _first_name(names.get(child_number) or names.get(str(child_number)))
    if child_name:
        return f"{child_name}'s {goal_suffix}"
    return f"Child {child_number}'s {goal_suffix}"


def display_cost_type(cost_type: str) -> str:
    return "Future" if "future" in cost_type.lower() else "Today"


def _rate_snapshot(calc) -> dict:
    if getattr(calc, "inflation_pre", None) is not None:
        return {
            "inflation_pre": calc.inflation_pre,
            "roi_pre": calc.roi_pre if calc.roi_pre is not None else 0.12,
            "inflation_post": calc.inflation_post if calc.inflation_post is not None else 0.06,
            "roi_post": calc.roi_post if calc.roi_post is not None else 0.08,
        }

    config = RateConfig.query.first()
    if config:
        return {
            "inflation_pre": config.inflation_pre,
            "roi_pre": config.roi_pre,
            "inflation_post": config.inflation_post,
            "roi_post": config.roi_post,
        }

    return {
        "inflation_pre": 0.06,
        "roi_pre": 0.12,
        "inflation_post": 0.06,
        "roi_post": 0.08,
    }


def _set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(
    cell,
    *,
    top: int = INSURANCE_CELL_MARGIN_V_TWIPS,
    bottom: int = INSURANCE_CELL_MARGIN_V_TWIPS,
    left: int = INSURANCE_CELL_MARGIN_H_TWIPS,
    right: int = INSURANCE_CELL_MARGIN_H_TWIPS,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr.findall(qn("w:tcMar"))):
        tc_pr.remove(child)
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        side_el = OxmlElement(f"w:{side}")
        side_el.set(qn("w:w"), str(value))
        side_el.set(qn("w:type"), "dxa")
        margins.append(side_el)
    tc_pr.append(margins)


def _set_row_height(row, height_twips: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for child in list(tr_pr.findall(qn("w:trHeight"))):
        tr_pr.remove(child)
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(tr_height)


def _apply_table_row_sizing(table) -> None:
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            height = INSURANCE_HEADER_ROW_HEIGHT_TWIPS
            font_pt = INSURANCE_HEADER_FONT_PT
            bold = True
        elif row_idx == len(table.rows) - 1:
            height = INSURANCE_FOOTER_ROW_HEIGHT_TWIPS
            font_pt = INSURANCE_HEADER_FONT_PT
            bold = True
        else:
            height = INSURANCE_DATA_ROW_HEIGHT_TWIPS
            font_pt = INSURANCE_BODY_FONT_PT
            bold = False

        _set_row_height(row, height)
        for cell in row.cells:
            _set_cell_margins(cell)
            align = cell.paragraphs[0].alignment
            _style_cell_text(cell, bold=bold, align=align, size_pt=font_pt)


def _style_cell_text(
    cell,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    size_pt: int = 9,
) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.25
    if not paragraph.runs:
        paragraph.add_run(cell.text)
    for run in paragraph.runs:
        run.bold = bold
        run.font.size = Pt(size_pt)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0, 0, 0)


def _emu_to_cm(value: int) -> float:
    return (value / EMU_PER_INCH) * 2.54


def _usable_page_width_cm(doc: Document) -> float:
    section = doc.sections[0]
    return section.page_width.cm - section.left_margin.cm - section.right_margin.cm


def _header_clearance_cm(doc: Document) -> float:
    body_xml = doc.element.body.xml
    anchors = re.findall(r"<wp:anchor[\s\S]*?</wp:anchor>", body_xml)
    if not anchors:
        return DEFAULT_TABLE_CLEARANCE_CM

    max_bottom_cm = 0.0
    for anchor in anchors:
        pos_v = re.search(r"<wp:positionV[\s\S]*?<wp:posOffset>(\d+)</wp:posOffset>", anchor)
        extent = re.search(r"<wp:extent cx=\"\d+\" cy=\"(\d+)\"", anchor)
        if not pos_v or not extent:
            continue
        bottom_cm = _emu_to_cm(int(pos_v.group(1)) + int(extent.group(1)))
        max_bottom_cm = max(max_bottom_cm, bottom_cm)

    if max_bottom_cm <= 0:
        return DEFAULT_TABLE_CLEARANCE_CM

    top_margin_cm = doc.sections[0].top_margin.cm
    clearance = max_bottom_cm - top_margin_cm + EXTRA_TABLE_GAP_CM
    return max(clearance, 0.0)


def _make_spacer_paragraph(before_cm: float):
    spacer = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(before_cm * TWIPS_PER_CM)))
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    spacer.append(p_pr)
    return spacer


def _apply_table_layout(doc: Document, table) -> None:
    usable_width_cm = _usable_page_width_cm(doc)
    ratio_total = sum(INSURANCE_COLUMN_RATIOS)
    col_widths = [
        Cm(usable_width_cm * ratio / ratio_total) for ratio in INSURANCE_COLUMN_RATIOS
    ]

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int(usable_width_cm * TWIPS_PER_CM)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tbl_pr.append(jc)


def _find_marker_paragraph(doc: Document):
    from app.services.report_service import _iter_xml_paragraphs, _xml_para_text

    for element in _iter_xml_paragraphs(doc):
        if INSURANCE_TABLE_MARKER in _xml_para_text(element):
            return element
    return None


def _clear_marker_paragraph(marker) -> None:
    from app.services.report_service import _set_xml_para_text

    _set_xml_para_text(marker, "")


def _insert_table_at_marker(doc: Document, marker, table) -> None:
    parent = marker.getparent()
    if parent is None:
        return

    body = doc.element.body
    if table._tbl.getparent() is not None:
        body.remove(table._tbl)

    insert_at = parent.index(marker)
    parent.remove(marker)

    clearance_cm = _header_clearance_cm(doc)
    if clearance_cm > 0:
        parent.insert(insert_at, _make_spacer_paragraph(clearance_cm))
        insert_at += 1

    parent.insert(insert_at, table._tbl)


def _set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)


def fill_insurance_table(
    doc: Document,
    rows: list[dict],
    total_pv: float,
    child_names_by_number: dict[int, str] | None = None,
) -> Document:
    """Replace {{insurance_table}} with a styled Word table."""
    marker = _find_marker_paragraph(doc)
    if marker is None:
        return doc

    # Skip zero-amount rows so empty retirement/household inputs are not printed.
    rows = [
        row
        for row in rows
        if float(row.get("Amount (₹)", 0) or 0) > 0
        or float(row.get("PV (₹)", 0) or 0) > 0
    ]

    row_count = 1 + len(rows) + 1
    table = doc.add_table(rows=row_count, cols=5)
    _set_table_borders(table)
    _apply_table_layout(doc, table)

    header_row = table.rows[0].cells
    for idx, label in enumerate(INSURANCE_HEADERS):
        header_row[idx].text = label
        _set_cell_shading(header_row[idx], INSURANCE_HEADER_FILL)
        _style_cell_text(
            header_row[idx],
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size_pt=INSURANCE_HEADER_FONT_PT,
        )

    for row_idx, row in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        values = [
            display_need_label(
                row.get("Need", ""),
                child_names_by_number=child_names_by_number,
            ),
            str(row.get("Years", "")),
            fmt_inr_table(float(row.get("Amount (₹)", 0))),
            display_cost_type(row.get("Type", "")),
            fmt_inr_table(float(row.get("PV (₹)", 0))),
        ]
        alignments = [
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ]
        for col_idx, (value, align) in enumerate(zip(values, alignments)):
            cells[col_idx].text = value
            _style_cell_text(
                cells[col_idx],
                align=align,
                size_pt=INSURANCE_BODY_FONT_PT,
            )

    footer_cells = table.rows[-1].cells
    footer_cells[0].text = "Total insurance need"
    _set_cell_shading(footer_cells[0], INSURANCE_HEADER_FILL)
    _style_cell_text(
        footer_cells[0],
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size_pt=INSURANCE_HEADER_FONT_PT,
    )

    merged_middle = footer_cells[1].merge(footer_cells[3])
    merged_middle.text = ""
    _set_cell_shading(merged_middle, INSURANCE_HEADER_FILL)

    footer_cells[4].text = fmt_inr_table(total_pv)
    _set_cell_shading(footer_cells[4], INSURANCE_HEADER_FILL)
    # Body size keeps long totals (e.g. ₹ 3,08,45,574) on one line with bold.
    _style_cell_text(
        footer_cells[4],
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size_pt=INSURANCE_BODY_FONT_PT,
    )

    _apply_table_row_sizing(table)

    _clear_marker_paragraph(marker)
    _insert_table_at_marker(doc, marker, table)
    return doc


def compute_insurance_rows_for_report(
    calc,
    personal,
    goals,
    children_map: dict | None = None,
) -> tuple[list[dict], float]:
    """Return insurance rows from stored snapshot or recompute from assessment data."""
    if getattr(calc, "insurance_items", None):
        return calc.insurance_items, float(calc.total_insurance_required or 0)

    from app.api.v1.calculate.routes import goals_to_formula_input

    children_map = children_map or {}
    rates = _rate_snapshot(calc)
    rr = real_rate(rates["roi_post"], rates["inflation_post"])
    monthly_eff_pre = monthly_effective_rate(rates["roi_pre"])

    household_monthly = float(getattr(calc, "household_monthly", None) or 0)
    client_annual = float(getattr(calc, "client_annual_ret_reqd", None) or 0)
    spouse_annual = float(getattr(calc, "spouse_annual_ret_reqd", None) or 0)

    client_years = max(
        0, (personal.client_retirement_age or 60) - (personal.client_age or 0)
    )
    has_spouse = bool(personal.spouse_name or personal.spouse_dob)
    if has_spouse and personal.spouse_age is not None:
        spouse_years = max(
            0, (personal.spouse_retirement_age or 55) - personal.spouse_age
        )
    else:
        spouse_years = 0
        spouse_annual = 0.0

    goal_results, _ = compute_all_goals(
        goals_to_formula_input(goals, children_map),
        monthly_eff_pre,
    )
    return build_insurance_rows(
        client_years,
        spouse_years,
        household_monthly,
        client_annual,
        spouse_annual,
        goal_results,
        rr,
        rates["roi_post"],
    )


def ensure_insurance_template(template_dir: Path) -> Path:
    """Create insurance_cal.docx with placeholder if it does not exist."""
    template_dir.mkdir(parents=True, exist_ok=True)
    path = template_dir / "insurance_cal.docx"
    if path.exists():
        return path

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Insurance Need Analysis")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Calibri"

    doc.add_paragraph()
    client_line = doc.add_paragraph("Client: {{client_name}}")
    client_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_line = doc.add_paragraph("Date: {{report_date}}")
    date_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()
    marker = doc.add_paragraph(INSURANCE_TABLE_MARKER)
    marker.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(path)
    return path
