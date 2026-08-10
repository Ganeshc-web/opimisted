import io
import json
import os
import platform
import re
import shutil
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional
from uuid import UUID

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from pypdf import PdfWriter

from app import db
from app.core.education_display import program_display_name
from app.core.formatters import fmt_inr_report, fmt_inr_report_blank
from app.core.formulas import excel_FV, LIFE_EXPECTANCY, PF_CONTRIBUTION_YEARLY_INCREASE
from app.core.goal_descriptions import (
    goal_description_from_picks,
    pick_goal_descriptions,
)
from app.models.education_db import EducationProgram
from app.models.family import FamilyDetails
from app.models.rate_config import RateConfig
from app.models.tour_db import TourDestination
from app.services.insurance_docx_service import (
    compute_insurance_rows_for_report,
    ensure_insurance_template,
    fill_insurance_table,
)
from app.services.summary_docx_service import (
    build_summary_rows,
    fill_summary_table,
    has_summary_content,
)

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
TEMPLATE_DIR = PROJECT_ROOT / "templates"
FLAGS_DIR = TEMPLATE_DIR / "flags"
FLAGS_SVG_DIR = FLAGS_DIR / "svg"
COUNTRIES_JSON_PATH = FLAGS_DIR / "countries.json"


def _children_for_assessment(assessment_id) -> list:
    aid = assessment_id if isinstance(assessment_id, UUID) else UUID(str(assessment_id))
    family = FamilyDetails.query.filter_by(assessment_id=aid).first()
    return list(family.children) if family else []


def _child_names_by_id(assessment_id: str) -> dict:
    return {c.id: _first_name(c.full_name) for c in _children_for_assessment(assessment_id)}


def _child_name_placeholders(assessment_id: str) -> dict[str, str]:
    """Placeholders for templates: child_1_name, child_2_name, ..."""
    out: dict[str, str] = {}
    for child in _children_for_assessment(assessment_id):
        out[f"child_{child.child_number}_name"] = _first_name(child.full_name)
    return out


def report_download_basename(personal: object | None) -> str:
    """Client-facing report name: UserName_GoalAnalysisReport."""
    raw = (getattr(personal, "client_name", None) or "Client").strip()
    safe = re.sub(r"[^\w\s\-]", "", raw, flags=re.UNICODE)
    safe = re.sub(r"[\s\-]+", "_", safe).strip("_")
    return f"{safe or 'Client'}_GoalAnalysisReport"


LIFESTYLE_TEMPLATE_MAP = {
    "Home Purchase": "house_purchase.docx",
    "Home Renovation": "house_renovation.docx",
    "Car Purchase": "car_purchase.docx",
    "Holiday Home": "holiday_home.docx",
    "Family Gifting": "family_gifting.docx",
    "Charity": "charity.docx",
    "Child Birth Expenses": "child_birth.docx",
    "Big Purchases": "big_purchases.docx",
    "Estate For Children": "estate_for_children.docx",
    "Foreign Tour": "foreign_tour.docx",
    "Other": "other_goals.docx",
}

CHILD_TEMPLATE_MAP = {
    "Graduation": "child1_graduation.docx",
    "Post Graduation": "child1_post_graduation.docx",
    "Marriage": "child1_marriage.docx",
    "Other": "child1_other.docx",
}

LIFESTYLE_GOAL_ORDER = list(LIFESTYLE_TEMPLATE_MAP.keys())
CHILD_GOAL_ORDER = list(CHILD_TEMPLATE_MAP.keys())
# Back-compat aliases used by older call sites / tests.
CHILD1_TEMPLATE_MAP = CHILD_TEMPLATE_MAP
CHILD1_GOAL_ORDER = CHILD_GOAL_ORDER

STATIC_TAIL_PAGES = [
    "insurance_cal.docx",
    "assume.docx",
    "services.docx",
    "testimonials.docx",
    "last_page.docx",
]

EDUCATION_RECOMMENDATION_LIMIT = 5
TOUR_BUDGET_TOLERANCE_PERCENT = 15
TOUR_RECOMMENDATION_LIMIT = 3
TOUR_COST_RANGE_FACTOR = 0.15


def reports_folder() -> str:
    """Writable folder for generated reports (production host disk or local dev)."""
    from flask import current_app, has_app_context

    if has_app_context():
        folder = current_app.config.get("REPORTS_FOLDER", "reports/")
    else:
        folder = os.environ.get("REPORTS_FOLDER", "reports/")
    folder = str(folder).strip()
    if folder.startswith(("/", "\\")) or (len(folder) > 1 and folder[1] == ":"):
        return folder
    return str(PROJECT_ROOT / folder)


def try_convert_docx_to_pdf(docx_path: str) -> Optional[str]:
    pdf_path = docx_path.replace(".docx", ".pdf")

    if platform.system() == "Windows":
        try:
            import pythoncom
            from docx2pdf import convert

            pythoncom.CoInitialize()
            try:
                convert(docx_path, pdf_path)
            finally:
                pythoncom.CoUninitialize()
            return pdf_path if os.path.exists(pdf_path) else None
        except Exception:
            return None

    try:
        out_dir = os.path.dirname(docx_path) or "."
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                out_dir,
                docx_path,
            ],
            check=True,
            capture_output=True,
            timeout=180,
        )
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception:
        pass

    return None


def _document_xml_roots(doc: Document):
    yield doc.element.body
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part:
                yield part._element


def _iter_xml_paragraphs(doc: Document):
    """Iterate all w:p elements, including those inside text boxes."""
    for root in _document_xml_roots(doc):
        for p in root.iter(qn("w:p")):
            yield p


def _is_inside_textbox(node) -> bool:
    """True if node is under a Word text-box content element."""
    cur = node.getparent() if hasattr(node, "getparent") else None
    while cur is not None:
        if cur.tag == qn("w:txbxContent"):
            return True
        cur = cur.getparent()
    return False


def _paragraph_t_nodes(p):
    """
    w:t nodes that belong to this paragraph's own text.

    Floating text boxes are nested under the anchor paragraph in Word XML.
    Recursive p.iter(w:t) would include those nested boxes and wipe them when
    the anchor paragraph is updated — so exclude any w:t inside txbxContent
    unless this paragraph itself lives inside a text box.
    """
    if _is_inside_textbox(p):
        return list(p.iter(qn("w:t")))

    nodes = []
    for t in p.iter(qn("w:t")):
        if _is_inside_textbox(t):
            continue
        nodes.append(t)
    return nodes


def _xml_para_text(p) -> str:
    return "".join(t.text or "" for t in _paragraph_t_nodes(p))


def _set_xml_para_text(p, new_text: str) -> None:
    t_nodes = _paragraph_t_nodes(p)
    if not t_nodes:
        return
    t_nodes[0].text = new_text
    for t in t_nodes[1:]:
        t.text = ""


def _goal_timeline_value(goal: dict) -> str:
    target_year = goal.get("target_year", "")
    goal_type = goal.get("goal_type", "")
    child_name = (goal.get("child_name") or "").strip()
    if target_year == "" and not goal_type:
        return ""
    if child_name:
        return f"{target_year} {child_name} {goal_type}".strip()
    return f"{target_year} {goal_type}".strip()


def _timeline_goals(goals_data: list[dict]) -> list[dict]:
    lifestyle, _ = _goal_lookup(goals_data)
    ordered: list[dict] = []
    for goal_type in LIFESTYLE_GOAL_ORDER:
        if goal_type in lifestyle:
            ordered.append(lifestyle[goal_type])
    # Custom Other goals are stored under their entered name, not "Other".
    for goal_type, goal in lifestyle.items():
        if goal_type not in LIFESTYLE_TEMPLATE_MAP:
            ordered.append(goal)
    # Include every child's goals (same templates reused for child 2+).
    ordered.extend(_iter_child_goals(goals_data))
    ordered.sort(key=lambda g: (g.get("target_year") is None, g.get("target_year") or 0))
    return ordered


def _textbox_vertical_pos(p) -> float:
    """Visual top position for a paragraph inside a floating text box (pt)."""
    node = p
    while node is not None:
        # DrawingML
        if node.tag == qn("wp:anchor"):
            pos_v = node.find(qn("wp:positionV"))
            if pos_v is not None:
                off = pos_v.find(qn("wp:posOffset"))
                if off is not None and off.text:
                    try:
                        return int(off.text) / 12700.0  # EMU → pt
                    except ValueError:
                        pass
        # VML
        if node.tag == "{urn:schemas-microsoft-com:vml}shape":
            style = node.get("style") or ""
            match = re.search(r"margin-top:([0-9.]+)pt", style)
            if match:
                return float(match.group(1))
        node = node.getparent()
    return 0.0


def replace_placeholders(
    doc: Document,
    replacements: dict,
    goals_data: list[dict] | None = None,
    timeline: bool = False,
    tour_slots: list[dict] | None = None,
) -> Document:
    timeline_goals = _timeline_goals(goals_data) if timeline and goals_data else []
    tour_name_idx = 0
    tour_cost_idx = 0
    tour_slots = tour_slots or []

    # Collect roadmap slots first and fill top→bottom (Word XML is often bottom→top).
    timeline_paras: list = []
    if timeline:
        for p in _iter_xml_paragraphs(doc):
            text = _xml_para_text(p)
            if "{{goals_timeline}}" not in text:
                continue
            if not _is_inside_textbox(p) and "}}{{" in text:
                _set_xml_para_text(p, "")
                continue
            if _is_inside_textbox(p):
                timeline_paras.append(p)
        timeline_paras.sort(key=_textbox_vertical_pos)
        for idx, p in enumerate(timeline_paras):
            goal = timeline_goals[idx] if idx < len(timeline_goals) else {}
            new_text = _xml_para_text(p).replace(
                "{{goals_timeline}}",
                _goal_timeline_value(goal),
                1,
            )
            _set_xml_para_text(p, new_text)

    for p in _iter_xml_paragraphs(doc):
        text = _xml_para_text(p)
        if "{{" not in text:
            continue

        # Word often keeps a shape-anchor dump paragraph that concatenates every
        # text-box placeholder (}}{{...}}). Clear it so values don't leak as
        # "5%5%1..." and so roadmap timeline slots aren't consumed twice.
        if not _is_inside_textbox(p) and "}}{{" in text:
            _set_xml_para_text(p, "")
            continue

        # Already handled above.
        if timeline and "{{goals_timeline}}" in text and _is_inside_textbox(p):
            continue

        new_text = text
        for key, value in replacements.items():
            new_text = new_text.replace(f"{{{{{key}}}}}", str(value))

        if tour_slots:
            while "{{tour_1_name}}" in new_text:
                slot = tour_slots[tour_name_idx] if tour_name_idx < len(tour_slots) else {}
                new_text = new_text.replace(
                    "{{tour_1_name}}",
                    slot.get("name", ""),
                    1,
                )
                tour_name_idx += 1
            while "{{tour_1_cost}}" in new_text:
                slot = tour_slots[tour_cost_idx] if tour_cost_idx < len(tour_slots) else {}
                new_text = new_text.replace(
                    "{{tour_1_cost}}",
                    slot.get("cost", ""),
                    1,
                )
                tour_cost_idx += 1

        if new_text != text:
            _set_xml_para_text(p, new_text)

    return doc


def _goal_lookup(goals_data: list[dict]) -> tuple[dict, dict]:
    """Return lifestyle goals and child-1 goals (legacy helper)."""
    lifestyle: dict[str, dict] = {}
    child1: dict[str, dict] = {}
    for goal in goals_data:
        if goal.get("category") == "child_goal":
            if goal.get("child_number") == 1:
                child1[goal["goal_type"]] = goal
        else:
            lifestyle[goal["goal_type"]] = goal
    return lifestyle, child1


def _iter_child_goals(goals_data: list[dict]) -> list[dict]:
    """All child goals ordered by child_number then goal type."""
    child_goals = [
        g for g in goals_data if g.get("category") == "child_goal" and g.get("goal_type")
    ]
    type_rank = {name: idx for idx, name in enumerate(CHILD_GOAL_ORDER)}

    def sort_key(goal: dict):
        number = goal.get("child_number")
        number_key = number if isinstance(number, int) else 10_000
        return (number_key, type_rank.get(goal.get("goal_type"), 99), goal.get("target_year") or 0)

    return sorted(child_goals, key=sort_key)


def _first_name(full_name: str | None) -> str:
    """Return only the first whitespace-delimited part of a child's full name."""
    parts = (full_name or "").strip().split()
    return parts[0] if parts else ""


def _possessive_name(name: str | None) -> str:
    """Format a first name for child goal templates, e.g. Aarav's."""
    first_name = _first_name(name)
    return f"{first_name}'s" if first_name else ""


def _goal_replacements(goal: dict, base: dict | None = None) -> dict:
    return {
        "goal_target_year": str(goal.get("target_year", "") or ""),
        "goal_current_cost": fmt_inr_report_blank(goal.get("today_cost")),
        "goal_future_cost": fmt_inr_report_blank(goal.get("future_cost")),
        "goal_monthly_sip": fmt_inr_report_blank(goal.get("monthly_sip")),
        "child_name": _possessive_name(goal.get("child_name")),
        "child_number": str(goal.get("child_number") or ""),
        "goal_description": goal_description_from_picks(
            goal.get("goal_type"), base
        ),
    }


def _education_goal_for_recommendations(goals_data: list[dict]) -> dict | None:
    for goal in _iter_child_goals(goals_data):
        if goal.get("goal_type") in ("Graduation", "Post Graduation"):
            return goal
    return None


def _education_programs_for_budget(
    budget: float,
    level: str | None = None,
    limit: int = EDUCATION_RECOMMENDATION_LIMIT,
) -> list[EducationProgram]:
    """Return programs at or under budget, closest first (no over-budget fallback)."""
    if budget <= 0:
        return []

    query = EducationProgram.query.filter(
        EducationProgram.approx_cost_inr > 0,
        EducationProgram.approx_cost_inr <= budget,
    )
    if level:
        query = query.filter(EducationProgram.level == level)

    return (
        query.order_by(db.func.abs(EducationProgram.approx_cost_inr - budget))
        .limit(limit)
        .all()
    )


def _education_replacements(goals_data: list[dict]) -> dict:
    replacements: dict[str, str] = {
        "edu_today_cost": "",
    }
    for i in range(1, EDUCATION_RECOMMENDATION_LIMIT + 1):
        replacements[f"edu_{i}_name"] = ""
        replacements[f"edu_{i}_cost"] = ""

    goal = _education_goal_for_recommendations(goals_data)
    if not goal:
        return replacements

    budget = float(goal.get("today_cost") or 0)
    replacements["edu_today_cost"] = fmt_inr_report(budget)

    level = goal.get("goal_type")
    if level not in ("Graduation", "Post Graduation"):
        level = None

    for index, program in enumerate(
        _education_programs_for_budget(budget, level=level),
        start=1,
    ):
        replacements[f"edu_{index}_name"] = program_display_name(program)
        replacements[f"edu_{index}_cost"] = fmt_inr_report(program.approx_cost_inr)

    return replacements


def _tour_goal_for_recommendations(goals_data: list[dict]) -> dict | None:
    lifestyle, _ = _goal_lookup(goals_data)
    return lifestyle.get("Foreign Tour")


def _tour_destinations_for_budget(
    budget: float,
    tolerance_percent: float = TOUR_BUDGET_TOLERANCE_PERCENT,
    limit: int = TOUR_RECOMMENDATION_LIMIT,
) -> list[TourDestination]:
    if budget <= 0:
        return []

    lower = budget * (1 - tolerance_percent / 100)
    upper = budget * (1 + tolerance_percent / 100)
    matches = (
        TourDestination.query.filter(
            TourDestination.budget_inr >= lower,
            TourDestination.budget_inr <= upper,
        )
        .order_by(db.func.abs(TourDestination.budget_inr - budget))
        .limit(limit)
        .all()
    )
    if matches:
        return matches

    # Fallback: nearest destinations by budget when none sit inside the band.
    return (
        TourDestination.query.order_by(
            db.func.abs(TourDestination.budget_inr - budget)
        )
        .limit(limit)
        .all()
    )


def _fmt_tour_budget_range(budget_inr: float) -> str:
    low = budget_inr * (1 - TOUR_COST_RANGE_FACTOR)
    high = budget_inr * (1 + TOUR_COST_RANGE_FACTOR)
    return f"{fmt_inr_report(low)} - {fmt_inr_report(high)}"


def _tour_slots(goals_data: list[dict]) -> list[dict]:
    slots: list[dict] = [
        {"name": "", "cost": ""} for _ in range(TOUR_RECOMMENDATION_LIMIT)
    ]

    goal = _tour_goal_for_recommendations(goals_data)
    if not goal:
        return slots

    budget = float(goal.get("today_cost") or 0)
    destinations = _tour_destinations_for_budget(budget)
    for index, destination in enumerate(destinations):
        slots[index] = {
            "name": destination.country or "",
            "cost": _fmt_tour_budget_range(destination.budget_inr or 0),
        }
    return slots


def _tour_replacements(goals_data: list[dict]) -> dict:
    replacements: dict[str, str] = {
        "tour_today_cost": "",
    }
    for i in range(1, TOUR_RECOMMENDATION_LIMIT + 1):
        replacements[f"tour_{i}_name"] = ""
        replacements[f"tour_{i}_cost"] = ""

    goal = _tour_goal_for_recommendations(goals_data)
    if goal:
        budget = float(goal.get("today_cost") or 0)
        replacements["tour_today_cost"] = fmt_inr_report(budget)

    slots = _tour_slots(goals_data)
    for index, slot in enumerate(slots, start=1):
        replacements[f"tour_{index}_name"] = slot["name"]
        replacements[f"tour_{index}_cost"] = slot["cost"]
    return replacements


FLAG_COUNTRY_ALIASES = {
    "bali (indonesia)": "ID",
    "dubai (uae)": "AE",
    "turkey": "TR",
    "south korea": "KR",
    "china": "CN",
}


def _country_flag_code(country_name: str) -> str | None:
    """Resolve a tour destination label to a lowercase ISO alpha-2 code."""
    normalized = (country_name or "").strip().casefold()
    if not normalized:
        return None
    if normalized in FLAG_COUNTRY_ALIASES:
        return FLAG_COUNTRY_ALIASES[normalized].lower()

    with COUNTRIES_JSON_PATH.open(encoding="utf-8") as handle:
        countries = json.load(handle)
    for code, name in countries.items():
        if str(name).strip().casefold() == normalized:
            return str(code).lower()
    return None


def _flag_png_bytes(country_name: str) -> bytes | None:
    code = _country_flag_code(country_name)
    if not code:
        return None
    svg_path = FLAGS_SVG_DIR / f"{code}.svg"
    if not svg_path.is_file():
        return None

    import resvg_py

    return resvg_py.svg_to_bytes(
        svg_path=str(svg_path),
        width=240,
        height=160,
    )


def _fill_tour_flags(doc: Document, tour_slots: list[dict]) -> Document:
    """Replace {{tour_N_flag}} table-cell placeholders with flag images."""
    slots_by_marker = {
        f"{{{{tour_{index}_flag}}}}": slot
        for index, slot in enumerate(tour_slots, start=1)
    }
    for element in _iter_xml_paragraphs(doc):
        marker = _xml_para_text(element).strip()
        if marker not in slots_by_marker:
            continue

        # Use the document proxy as parent so python-docx can access its part
        # when creating the image relationship.
        paragraph = Paragraph(element, doc)
        paragraph.clear()

        # Top-align the flag so it starts on the same row as the tour name.
        parent = element.getparent()
        while parent is not None and parent.tag != qn("w:tc"):
            parent = parent.getparent()
        if parent is not None:
            tc_pr = parent.get_or_add_tcPr()
            v_align = tc_pr.find(qn("w:vAlign"))
            if v_align is None:
                v_align = OxmlElement("w:vAlign")
                tc_pr.append(v_align)
            v_align.set(qn("w:val"), "top")

        country_name = slots_by_marker[marker].get("name", "")
        image_bytes = _flag_png_bytes(country_name)
        if image_bytes:
            paragraph.add_run().add_picture(
                io.BytesIO(image_bytes),
                width=Inches(1.0),
            )
    return doc


def has_retirement_planning(calc) -> bool:
    """True when the user supplied retirement-planning inputs (or meaningful outputs)."""
    if calc is None:
        return False
    attrs = (
        "household_monthly",
        "client_annual_ret_reqd",
        "spouse_annual_ret_reqd",
        "client_corpus",
        "client_pf_corpus",
        "client_provisions_made",
        "client_monthly_sip",
        "client_lump_sum",
        "spouse_corpus",
        "spouse_pf_corpus",
        "spouse_provisions_made",
        "spouse_monthly_sip",
        "spouse_lump_sum",
    )
    for attr in attrs:
        try:
            if float(getattr(calc, attr, None) or 0) > 0:
                return True
        except (TypeError, ValueError):
            continue
    return False


def build_page_list(
    goals_data: list[dict],
    *,
    include_retirement: bool = True,
    include_summary: bool = True,
) -> list[dict]:
    """
    Build ordered report pages.

    Each item is {"template": "<file.docx>", "goal": <goal dict|None>}.
    Child goal pages reuse the same child1_*.docx templates for every child.
    Retirement/summary pages are omitted when their include_* flags are False.
    Summary sits after insurance_cal.docx and before assume.docx.
    """
    pages: list[dict] = [
        {"template": "cover.docx", "goal": None},
        {"template": "roadmap.docx", "goal": None},
    ]
    lifestyle, _ = _goal_lookup(goals_data)

    for goal_type in LIFESTYLE_GOAL_ORDER:
        if goal_type in lifestyle:
            pages.append(
                {"template": LIFESTYLE_TEMPLATE_MAP[goal_type], "goal": lifestyle[goal_type]}
            )
            if goal_type == "Foreign Tour":
                pages.append({"template": "suggested_foreign.docx", "goal": None})

    # Custom Other goals saved under the entered name use the Other template.
    for goal_type, goal in lifestyle.items():
        if goal_type not in LIFESTYLE_TEMPLATE_MAP:
            pages.append(
                {"template": LIFESTYLE_TEMPLATE_MAP["Other"], "goal": goal}
            )

    child_goals = _iter_child_goals(goals_data)
    for goal in child_goals:
        template = CHILD_TEMPLATE_MAP.get(goal["goal_type"]) or CHILD_TEMPLATE_MAP["Other"]
        pages.append({"template": template, "goal": goal})

    if any(g.get("goal_type") in ("Graduation", "Post Graduation") for g in child_goals):
        pages.append({"template": "suggested_uni_coll.docx", "goal": None})

    if include_retirement:
        pages.append({"template": "retirement_planning.docx", "goal": None})
    for template in STATIC_TAIL_PAGES:
        pages.append({"template": template, "goal": None})
        if template == "insurance_cal.docx" and include_summary:
            pages.append({"template": "summary.docx", "goal": None})
    return pages


def _strip_life_expectancy_years_suffix(doc: Document, age_value: str) -> Document:
    """
    Keep life expectancy as digits only in the value box.

    Some templates still contain '{{life_expectancy}} years'; after fill that
    becomes '80 years' and overflows LibreOffice text boxes. Strip the suffix.
    """
    age = str(age_value or "").strip()
    if not age:
        return doc

    targets = (
        f"{age} years",
        f"{age} Years",
        f"{age} YEARS",
    )
    for p in _iter_xml_paragraphs(doc):
        text = _xml_para_text(p)
        if not text:
            continue
        new_text = text
        for target in targets:
            if target in new_text:
                new_text = new_text.replace(target, age)
        if new_text != text:
            _set_xml_para_text(p, new_text)
    return doc


def fill_template(
    template_name: str,
    output_path: str,
    replacements: dict,
    goals_data: list[dict] | None = None,
    *,
    goal: dict | None = None,
    calc=None,
    personal=None,
    goals_orm=None,
    children_map: dict | None = None,
) -> str:
    ensure_insurance_template(TEMPLATE_DIR)
    doc = Document(str(TEMPLATE_DIR / template_name))

    if template_name == "roadmap.docx":
        doc = replace_placeholders(doc, replacements, goals_data=goals_data, timeline=True)
    elif template_name == "assume.docx":
        doc = replace_placeholders(doc, replacements)
        doc = _strip_life_expectancy_years_suffix(
            doc, str(replacements.get("life_expectancy") or "")
        )
    elif template_name == "suggested_uni_coll.docx":
        doc = replace_placeholders(
            doc,
            {**replacements, **_education_replacements(goals_data or [])},
        )
    elif template_name == "suggested_foreign.docx":
        tour_slots = _tour_slots(goals_data or [])
        doc = _fill_tour_flags(doc, tour_slots)
        doc = replace_placeholders(
            doc,
            {**replacements, **_tour_replacements(goals_data or [])},
        )
    elif template_name == "insurance_cal.docx":
        doc = replace_placeholders(doc, replacements)
        if calc is not None and personal is not None and goals_orm is not None:
            insurance_rows, insurance_total = compute_insurance_rows_for_report(
                calc, personal, goals_orm, children_map
            )
            child_names_by_number: dict[int, str] = {}
            assessment_id = getattr(calc, "assessment_id", None)
            if assessment_id is not None:
                child_names_by_number = {
                    int(c.child_number): _first_name(c.full_name)
                    for c in _children_for_assessment(assessment_id)
                    if c.child_number is not None and (c.full_name or "").strip()
                }
            if not child_names_by_number and goals_data:
                for g in goals_data:
                    num = g.get("child_number")
                    name = (g.get("child_name") or "").strip()
                    if num is not None and name:
                        child_names_by_number[int(num)] = _first_name(name)
            doc = fill_insurance_table(
                doc,
                insurance_rows,
                insurance_total,
                child_names_by_number=child_names_by_number,
            )
        else:
            from app.services.insurance_docx_service import (
                _clear_marker_paragraph,
                _find_marker_paragraph,
            )

            marker = _find_marker_paragraph(doc)
            if marker is not None:
                _clear_marker_paragraph(marker)
    elif template_name == "summary.docx":
        doc = replace_placeholders(doc, replacements)
        summary_rows, summary_total = build_summary_rows(
            goals_data or [],
            calc=calc,
            personal=personal,
        )
        if summary_rows:
            doc = fill_summary_table(doc, summary_rows, summary_total)
        else:
            from app.services.summary_docx_service import (
                _clear_marker_paragraph as _clear_summary_marker,
                _find_marker_paragraph as _find_summary_marker,
            )

            marker = _find_summary_marker(doc)
            if marker is not None:
                _clear_summary_marker(marker)
    elif goal is not None:
        doc = replace_placeholders(
            doc,
            {**replacements, **_goal_replacements(goal, replacements)},
        )
    else:
        doc = replace_placeholders(doc, replacements)

    doc.save(output_path)
    return output_path


def merge_pdfs(pdf_paths: list[str], output_path: str) -> str:
    writer = PdfWriter()
    for path in pdf_paths:
        writer.append(path)
    with open(output_path, "wb") as f:
        writer.write(f)
    writer.close()
    return output_path


def _rate_snapshot(calc: object) -> dict:
    if getattr(calc, "inflation_pre", None) is not None:
        return {
            "inflation_pre": calc.inflation_pre,
            "roi_pre": calc.roi_pre if calc.roi_pre is not None else 0.12,
            "inflation_post": calc.inflation_post if calc.inflation_post is not None else 0.06,
            "roi_post": calc.roi_post if calc.roi_post is not None else 0.08,
        }

    config = RateConfig.query.first()
    if config:
        return {
            "inflation_pre": config.inflation_pre,
            "roi_pre": config.roi_pre,
            "inflation_post": config.inflation_post,
            "roi_post": config.roi_post,
        }

    return {
        "inflation_pre": 0.06,
        "roi_pre": 0.12,
        "inflation_post": 0.06,
        "roi_post": 0.08,
    }


def _years_to_retirement(age, retirement_age) -> int:
    if age is None or retirement_age is None:
        return 0
    return max(retirement_age - age, 0)


def _expense_placeholders(annual_ret_reqd, years_to_retirement, inflation_pre) -> tuple[str, str]:
    if not annual_ret_reqd or years_to_retirement is None:
        return "", ""
    expense_today_pm = annual_ret_reqd / 12
    expense_at_ret_pm = excel_FV(inflation_pre, years_to_retirement, 0, expense_today_pm)
    return (
        fmt_inr_report_blank(expense_today_pm),
        fmt_inr_report_blank(expense_at_ret_pm),
    )


def _build_replacements(
    assessment_id: str,
    calc: object,
    personal: object,
    comm: object,
) -> dict:
    rates = _rate_snapshot(calc)

    client_ret_age = personal.client_retirement_age or 60
    client_age = personal.client_age
    client_yrs_to_ret = (
        _years_to_retirement(client_age, client_ret_age) if client_age is not None else 0
    )
    client_ret_period = LIFE_EXPECTANCY - client_ret_age

    client_exp_today, client_exp_at_ret = _expense_placeholders(
        getattr(calc, "client_annual_ret_reqd", None) or 0,
        client_yrs_to_ret,
        rates["inflation_pre"],
    )

    spouse_ret_age = personal.spouse_retirement_age or 55
    spouse_age = personal.spouse_age
    if spouse_age is not None:
        spouse_yrs_to_ret = _years_to_retirement(spouse_age, spouse_ret_age)
        spouse_exp_today, spouse_exp_at_ret = _expense_placeholders(
            getattr(calc, "spouse_annual_ret_reqd", None) or 0,
            spouse_yrs_to_ret,
            rates["inflation_pre"],
        )
    else:
        spouse_yrs_to_ret = 0
        spouse_exp_today = ""
        spouse_exp_at_ret = ""

    has_spouse = bool(personal.spouse_name or personal.spouse_dob)

    replacements = {
        "client_name": personal.client_name or "",
        "client_age": str(personal.client_age or ""),
        "client_occupation": personal.client_occupation or "",
        "client_company": personal.client_company or "",
        "client_dob": str(personal.client_dob or ""),
        "client_retirement_age": str(client_ret_age),
        "client_retirement_period": str(client_ret_period),
        "client_years_to_retirement": str(client_yrs_to_ret) if client_yrs_to_ret else "",
        "client_expense_today": client_exp_today,
        "client_expense_at_retirement": client_exp_at_ret,
        "spouse_name": personal.spouse_name or "",
        "spouse_age": str(personal.spouse_age or ""),
        "spouse_occupation": personal.spouse_occupation or "",
        "spouse_company": personal.spouse_company or "",
        "spouse_dob": str(personal.spouse_dob or ""),
        "spouse_retirement_age": str(spouse_ret_age) if has_spouse else "",
        "spouse_years_to_retirement": str(spouse_yrs_to_ret) if has_spouse and spouse_yrs_to_ret else "",
        "spouse_expense_today": spouse_exp_today,
        "spouse_expense_at_retirement": spouse_exp_at_ret,
        "mobile": comm.mobile or "",
        "email": comm.email or "",
        "residential_address": comm.residential_address or "",
        "report_date": datetime.now(timezone.utc).strftime("%d %B %Y %H:%M:%S"),
        " report_date": datetime.now(timezone.utc).strftime("%d %B %Y %H:%M:%S"),
        "assessment_id": str(assessment_id),
        "client_corpus": fmt_inr_report_blank(calc.client_corpus),
        "client_pf_corpus": fmt_inr_report_blank(calc.client_pf_corpus),
        "client_provisions_made": fmt_inr_report_blank(
            getattr(calc, "client_provisions_made", None)
        ),
        "client_net_corpus": fmt_inr_report_blank(calc.client_net_corpus),
        "client_monthly_sip": fmt_inr_report_blank(calc.client_monthly_sip),
        "client_lump_sum": fmt_inr_report_blank(calc.client_lump_sum),
        "spouse_corpus": fmt_inr_report_blank(calc.spouse_corpus),
        "spouse_pf_corpus": fmt_inr_report_blank(calc.spouse_pf_corpus),
        "spouse_provisions_made": fmt_inr_report_blank(
            getattr(calc, "spouse_provisions_made", None)
        ),
        "spouse_net_corpus": fmt_inr_report_blank(calc.spouse_net_corpus),
        "spouse_monthly_sip": fmt_inr_report_blank(calc.spouse_monthly_sip),
        "spouse_lump_sum": fmt_inr_report_blank(calc.spouse_lump_sum),
        "total_insurance_required": fmt_inr_report_blank(calc.total_insurance_required),
        "total_goals_monthly_sip": fmt_inr_report_blank(calc.total_goals_monthly_sip),
        "inflation_pre": f"{rates['inflation_pre'] * 100:.1f}%",
        "roi_pre": f"{rates['roi_pre'] * 100:.1f}%",
        "inflation_post": f"{rates['inflation_post'] * 100:.1f}%",
        "roi_post": f"{rates['roi_post'] * 100:.1f}%",
        "life_expectancy": str(int(LIFE_EXPECTANCY)),
        "pf_contribution_yearly_increase": f"{PF_CONTRIBUTION_YEARLY_INCREASE * 100:.0f}%",
        "calculated_at": (
            calc.calculated_at.strftime("%d %B %Y %H:%M") if calc.calculated_at else ""
        ),
    }
    replacements.update(pick_goal_descriptions(seed=str(assessment_id)))
    return replacements


def generate_report(
    assessment_id: str,
    calc: object,
    personal: object,
    comm: object,
    goals: list,
    children_map: dict | None = None,
    convert_pdf: bool = True,
) -> dict:
    """
    Generate a modular per-page DOCX report, convert each page to PDF, and merge.
    """
    folder = reports_folder()
    os.makedirs(folder, exist_ok=True)

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    download_stem = report_download_basename(personal)
    # Unique on-disk stem; client-facing attach_name stays UserName_GoalAnalysisReport.pdf
    storage_stem = f"{download_stem}_{timestamp}"
    parts_dir = os.path.join(folder, storage_stem, "_parts")
    os.makedirs(parts_dir, exist_ok=True)

    children_map = children_map or {}
    replacements = _build_replacements(assessment_id, calc, personal, comm)
    replacements.setdefault("child_name", "")
    replacements.setdefault("child_number", "")
    replacements.update(_child_name_placeholders(assessment_id))

    names_by_id = _child_names_by_id(assessment_id)
    goals_data = [
        {
            "category": g.category,
            "goal_type": g.goal_type,
            "child_id": getattr(g, "child_id", None),
            "child_number": children_map.get(getattr(g, "child_id", None)),
            "child_name": names_by_id.get(getattr(g, "child_id", None), ""),
            "target_year": g.target_year,
            "today_cost": g.today_cost or 0,
            "future_cost": g.future_cost or 0,
            "monthly_sip": g.monthly_sip or 0,
        }
        for g in goals
    ]

    pages = build_page_list(
        goals_data,
        include_retirement=has_retirement_planning(calc),
        include_summary=has_summary_content(goals_data, calc),
    )
    docx_paths: list[str] = []

    for i, page in enumerate(pages):
        template_name = page["template"]
        goal = page.get("goal")
        # Unique part filenames when the same child template is reused.
        suffix = ""
        if goal and goal.get("category") == "child_goal":
            suffix = f"_c{goal.get('child_number') or i}"
        docx_out = os.path.join(parts_dir, f"{i:02d}_{Path(template_name).stem}{suffix}.docx")
        fill_template(
            template_name,
            docx_out,
            replacements,
            goals_data,
            goal=goal,
            calc=calc,
            personal=personal,
            goals_orm=goals,
            children_map=children_map,
        )
        docx_paths.append(docx_out)

    pdf_path = None
    if convert_pdf:
        pdf_parts: list[str] = []
        for docx_path in docx_paths:
            converted = try_convert_docx_to_pdf(docx_path)
            if not converted:
                raise RuntimeError(f"Failed to convert {docx_path} to PDF")
            pdf_parts.append(converted)
        pdf_path = os.path.join(folder, f"{storage_stem}.pdf")
        merge_pdfs(pdf_parts, pdf_path)

    shutil.rmtree(parts_dir, ignore_errors=True)

    attach_name = (
        f"{download_stem}.pdf" if convert_pdf and pdf_path else f"{download_stem}.docx"
    )
    result = {
        "file_name": attach_name,
        "docx_path": None,
        "pdf_path": pdf_path,
        "pages": [p["template"] for p in pages],
        "attach_path": pdf_path,
        "attach_name": attach_name,
    }

    from app.services.s3_storage import persist_report_files

    return persist_report_files(result, str(assessment_id))
