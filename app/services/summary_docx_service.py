"""Build the dynamic investment summary table inside summary.docx."""

from __future__ import annotations

from datetime import datetime

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.services.insurance_docx_service import (
    EXTRA_TABLE_GAP_CM,
    INSURANCE_HEADER_FILL,
    TWIPS_PER_CM,
    _emu_to_cm,
    _header_clearance_cm,
    _make_spacer_paragraph,
    _set_cell_margins,
    _set_cell_shading,
    _set_row_height,
    _set_table_borders,
    _usable_page_width_cm,
    fmt_inr_table,
)

SUMMARY_TABLE_MARKER = "{{summary_table}}"
SUMMARY_HEADERS = ["Goals", "Target Year", "Monthly Investment"]
SUMMARY_COLUMN_RATIOS = (5.0, 2.2, 3.0)
# Keep a little air under the decorative header image / title band.
SUMMARY_MIN_TOP_GAP_CM = 0.35
# Compact sizing so header + many goals + total stay on one A4 page.
SUMMARY_HEADER_FONT_PT = 11
SUMMARY_BODY_FONT_PT = 10
SUMMARY_HEADER_ROW_HEIGHT_TWIPS = 480
SUMMARY_DATA_ROW_HEIGHT_TWIPS = 400
SUMMARY_FOOTER_ROW_HEIGHT_TWIPS = 460
SUMMARY_CELL_MARGIN_V_TWIPS = 60
SUMMARY_CELL_MARGIN_H_TWIPS = 80
SUMMARY_MAX_TOP_CLEARANCE_CM = 5.8


def _find_marker_paragraph(doc: Document):
    from app.services.report_service import _iter_xml_paragraphs, _xml_para_text

    for element in _iter_xml_paragraphs(doc):
        if SUMMARY_TABLE_MARKER in _xml_para_text(element):
            return element
    return None


def _clear_marker_paragraph(marker) -> None:
    from app.services.report_service import _set_xml_para_text

    _set_xml_para_text(marker, "")


def _insert_table_at_marker(doc: Document, marker, table) -> None:
    parent = marker.getparent()
    if parent is None:
        return

    body = doc.element.body
    if table._tbl.getparent() is not None:
        body.remove(table._tbl)

    insert_at = parent.index(marker)
    parent.remove(marker)

    preceding_inline_height = 0.0
    for child in list(parent)[:insert_at]:
        for extent in child.iter(qn("wp:extent")):
            cy = extent.get("cy")
            if cy:
                try:
                    preceding_inline_height += _emu_to_cm(int(cy))
                except ValueError:
                    pass

    if "<wp:anchor" in doc.element.body.xml:
        clearance_cm = min(_header_clearance_cm(doc), SUMMARY_MAX_TOP_CLEARANCE_CM)
    elif preceding_inline_height > 0:
        # Decorative header image already consumed vertical space.
        clearance_cm = EXTRA_TABLE_GAP_CM
    else:
        clearance_cm = SUMMARY_MIN_TOP_GAP_CM

    if clearance_cm > 0:
        parent.insert(insert_at, _make_spacer_paragraph(clearance_cm))
        insert_at += 1

    parent.insert(insert_at, table._tbl)


def _style_summary_cell_text(
    cell,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    size_pt: int = SUMMARY_BODY_FONT_PT,
) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    if not paragraph.runs:
        paragraph.add_run(cell.text)
    for run in paragraph.runs:
        run.bold = bold
        run.font.size = Pt(size_pt)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Prefer keeping the table together across page breaks.
    p_pr = paragraph._p.get_or_add_pPr()
    for tag in ("w:keepNext", "w:keepLines"):
        existing = p_pr.find(qn(tag))
        if existing is not None:
            p_pr.remove(existing)
        p_pr.append(OxmlElement(tag))


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:cantSplit"))
    if existing is not None:
        tr_pr.remove(existing)
    tr_pr.append(OxmlElement("w:cantSplit"))


def _apply_summary_row_sizing(table) -> None:
    last_idx = len(table.rows) - 1
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            height = SUMMARY_HEADER_ROW_HEIGHT_TWIPS
            font_pt = SUMMARY_HEADER_FONT_PT
            bold = True
        elif row_idx == last_idx:
            height = SUMMARY_FOOTER_ROW_HEIGHT_TWIPS
            font_pt = SUMMARY_HEADER_FONT_PT
            bold = True
        else:
            height = SUMMARY_DATA_ROW_HEIGHT_TWIPS
            font_pt = SUMMARY_BODY_FONT_PT
            bold = False

        _set_row_height(row, height)
        # Exact heights keep Word/LibreOffice from growing rows past the budget.
        tr_height = row._tr.trPr.find(qn("w:trHeight"))
        if tr_height is not None:
            tr_height.set(qn("w:hRule"), "exact")
        _set_row_cant_split(row)

        for cell in row.cells:
            _set_cell_margins(
                cell,
                top=SUMMARY_CELL_MARGIN_V_TWIPS,
                bottom=SUMMARY_CELL_MARGIN_V_TWIPS,
                left=SUMMARY_CELL_MARGIN_H_TWIPS,
                right=SUMMARY_CELL_MARGIN_H_TWIPS,
            )
            align = cell.paragraphs[0].alignment or WD_ALIGN_PARAGRAPH.CENTER
            _style_summary_cell_text(
                cell,
                bold=bold,
                align=align,
                size_pt=font_pt,
            )
            # Last row should not force a following orphan paragraph.
            if row_idx == last_idx:
                p_pr = cell.paragraphs[0]._p.get_or_add_pPr()
                keep_next = p_pr.find(qn("w:keepNext"))
                if keep_next is not None:
                    p_pr.remove(keep_next)


def _apply_summary_table_layout(doc: Document, table) -> None:
    usable_width_cm = _usable_page_width_cm(doc)
    # Slightly inset so the centered table reads clearly inside wide margins.
    table_width_cm = max(usable_width_cm * 0.92, usable_width_cm - 1.0)
    ratio_total = sum(SUMMARY_COLUMN_RATIOS)
    col_widths = [
        Cm(table_width_cm * ratio / ratio_total) for ratio in SUMMARY_COLUMN_RATIOS
    ]

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    # Remove python-docx defaults that confuse LibreOffice (auto width / dup jc).
    for child in list(tbl_pr):
        if child.tag in (qn("w:tblW"), qn("w:jc"), qn("w:tblLayout")):
            tbl_pr.remove(child)

    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int(table_width_cm * TWIPS_PER_CM)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tbl_pr.append(jc)



def _summary_goal_label(goal: dict) -> str:
    goal_type = (goal.get("goal_type") or "").strip()
    if goal.get("category") == "child_goal":
        child_name = (goal.get("child_name") or "").strip()
        if child_name:
            return f"{child_name}'s {goal_type}".strip()
        child_number = goal.get("child_number")
        if child_number is not None:
            return f"Child {child_number}'s {goal_type}".strip()
    return goal_type


def _monthly_amount(value) -> float:
    """Monthly investment as a positive amount.

    Retirement SIPs come from Excel-style PMT and are stored negative, while goal
    SIPs are positive; every display path shows the magnitude.
    """
    try:
        return abs(float(value or 0))
    except (TypeError, ValueError):
        return 0.0


def _retirement_target_year(age, retirement_age, dob=None) -> str:
    if dob is not None and retirement_age is not None:
        try:
            return str(int(dob.year) + int(retirement_age))
        except (TypeError, ValueError, AttributeError):
            pass
    if age is None or retirement_age is None:
        return ""
    try:
        years = max(int(retirement_age) - int(age), 0)
    except (TypeError, ValueError):
        return ""
    return str(datetime.now().year + years)


def build_summary_rows(
    goals_data: list[dict],
    calc=None,
    personal=None,
) -> tuple[list[dict], float]:
    """
    Build summary rows for lifestyle/child goals plus retirement SIPs.

    Each row: {"goal": str, "target_year": str|int, "monthly_sip": float,
    "is_retirement": bool}
    """
    from app.services.report_service import _timeline_goals

    rows: list[dict] = []
    for goal in _timeline_goals(goals_data or []):
        sip = _monthly_amount(goal.get("monthly_sip"))
        if sip <= 0:
            continue
        rows.append(
            {
                "goal": _summary_goal_label(goal),
                "target_year": goal.get("target_year") or "",
                "monthly_sip": sip,
                "is_retirement": False,
            }
        )

    if calc is not None:
        client_sip = _monthly_amount(getattr(calc, "client_monthly_sip", None))
        if client_sip > 0:
            ret_age = getattr(personal, "client_retirement_age", None) if personal else None
            age = getattr(personal, "client_age", None) if personal else None
            dob = getattr(personal, "client_dob", None) if personal else None
            rows.append(
                {
                    "goal": "Retirement Planning",
                    "target_year": _retirement_target_year(age, ret_age, dob),
                    "monthly_sip": client_sip,
                    "is_retirement": True,
                }
            )

        spouse_sip = _monthly_amount(getattr(calc, "spouse_monthly_sip", None))
        if spouse_sip > 0:
            ret_age = getattr(personal, "spouse_retirement_age", None) if personal else None
            age = getattr(personal, "spouse_age", None) if personal else None
            dob = getattr(personal, "spouse_dob", None) if personal else None
            spouse_name = ""
            if personal is not None:
                parts = (getattr(personal, "spouse_name", None) or "").strip().split()
                spouse_name = parts[0] if parts else ""
            label = f"{spouse_name}'s Retirement" if spouse_name else "Spouse Retirement Planning"
            rows.append(
                {
                    "goal": label,
                    "target_year": _retirement_target_year(age, ret_age, dob),
                    "monthly_sip": spouse_sip,
                    "is_retirement": True,
                }
            )

    total = sum(_monthly_amount(r.get("monthly_sip")) for r in rows)
    return rows, total


def has_summary_content(goals_data: list[dict] | None = None, calc=None) -> bool:
    """True when at least one positive monthly investment exists for the summary page."""
    rows, total = build_summary_rows(goals_data or [], calc=calc, personal=None)
    if total > 0:
        return True
    # build_summary_rows without personal still includes retirement SIPs from calc
    if calc is not None:
        for attr in ("client_monthly_sip", "spouse_monthly_sip"):
            if _monthly_amount(getattr(calc, attr, None)) > 0:
                return True
    for goal in goals_data or []:
        if _monthly_amount(goal.get("monthly_sip")) > 0:
            return True
    return False


def fill_summary_table(
    doc: Document,
    rows: list[dict],
    total_monthly: float,
) -> Document:
    """Replace {{summary_table}} with a styled Word table."""
    marker = _find_marker_paragraph(doc)
    if marker is None:
        return doc

    rows = [row for row in rows if _monthly_amount(row.get("monthly_sip")) > 0]
    row_count = 1 + len(rows) + 1
    table = doc.add_table(rows=row_count, cols=3)
    _set_table_borders(table)
    _apply_summary_table_layout(doc, table)

    header_row = table.rows[0].cells
    for idx, label in enumerate(SUMMARY_HEADERS):
        header_row[idx].text = label
        _set_cell_shading(header_row[idx], INSURANCE_HEADER_FILL)
        _style_summary_cell_text(
            header_row[idx],
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size_pt=SUMMARY_HEADER_FONT_PT,
        )

    for row_idx, row in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        values = [
            str(row.get("goal") or ""),
            str(row.get("target_year") or ""),
            fmt_inr_table(_monthly_amount(row.get("monthly_sip"))),
        ]
        alignments = [
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ]
        for col_idx, (value, align) in enumerate(zip(values, alignments)):
            cells[col_idx].text = value
            _style_summary_cell_text(
                cells[col_idx],
                align=align,
                size_pt=SUMMARY_BODY_FONT_PT,
            )

    footer_cells = table.rows[-1].cells
    footer_cells[0].text = "Total Monthly Investment"
    _set_cell_shading(footer_cells[0], INSURANCE_HEADER_FILL)
    _style_summary_cell_text(
        footer_cells[0],
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size_pt=SUMMARY_HEADER_FONT_PT,
    )

    footer_cells[1].text = ""
    _set_cell_shading(footer_cells[1], INSURANCE_HEADER_FILL)

    footer_cells[2].text = fmt_inr_table(total_monthly)
    _set_cell_shading(footer_cells[2], INSURANCE_HEADER_FILL)
    _style_summary_cell_text(
        footer_cells[2],
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size_pt=SUMMARY_HEADER_FONT_PT,
    )

    _apply_summary_row_sizing(table)
    _clear_marker_paragraph(marker)
    _insert_table_at_marker(doc, marker, table)
    return doc
